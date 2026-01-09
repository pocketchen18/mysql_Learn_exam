import docx
import os

def extract_docx(file_path, output_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    doc = docx.Document(file_path)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    
    # Also extract from tables if any
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            content.append("\t".join(row_text))

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(content))
    print(f"Extracted {file_path} to {output_path}")

if __name__ == "__main__":
    base_dir = r'd:\code\Python\mysql_Learn_exam\base'
    extract_docx(os.path.join(base_dir, '测试题.docx'), 'test_questions.txt')
    extract_docx(os.path.join(base_dir, '数据库题目大集合.docx'), 'all_questions.txt')
